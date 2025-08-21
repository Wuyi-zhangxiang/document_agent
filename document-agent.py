
import os
import tempfile
import json5
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.text.paragraph import Paragraph
from docx.table import Table
from bs4 import BeautifulSoup
import re
from qwen_agent.log import get_edit_count

from qwen_agent.agents import VirtualMemoryAgent
from qwen_agent.gui import WebUI
from qwen_agent.tools.base import BaseTool, register_tool

@register_tool('markdown_splitter')
class MarkdownSplitter(BaseTool):
    description = '按H1标题(#)拆分Markdown文档为多个章节文件。输入: Markdown文件路径。输出: 包含各章节文件的列表，每个文件以序号和章节标题命名'
    parameters = [{
        'name': 'file_path',
        'type': 'string',
        'description': '待拆分的Markdown文件路径',
        'required': True
    }]

    def call(self, params: str, **kwargs) -> str:
        try:
            data = json5.loads(params)
            if 'file_path' not in data:
                raise ValueError("缺少必要参数: file_path")
            return self._split_markdown(data['file_path'])
        except Exception as e:
            return json5.dumps({
                'error': str(e),
                'status': 'failed'
            })

    def _split_markdown(self, md_path):
        """按H1标题拆分Markdown文件"""
        try:
            # 确保文件存在
            if not os.path.exists(md_path):
                raise FileNotFoundError(f"文件不存在: {md_path}")
            
            # 读取文件内容
            with open(md_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 使用改进的正则表达式拆分
            sections = re.split(r'^#\s+(.+)$', content, flags=re.MULTILINE)
            output_files = []
            
            # 处理每个章节
            for i in range(1, len(sections), 2):
                title = sections[i].strip()
                content = sections[i+1] if i+1 < len(sections) else ""
                
                # 生成更安全的文件名
                safe_title = re.sub(r'[^\w\-_\. ]', '_', title)
                filename = f"{i//2+1:02d}-{safe_title}.md"
                output_path = os.path.join(os.path.dirname(md_path), filename)
                
                # 写入文件并验证
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(f"# {title}\n\n{content}")
                    f.flush()
                
                if not os.path.exists(output_path):
                    raise IOError(f"文件创建失败: {output_path}")
                
                output_files.append(output_path)
            
            return json5.dumps({
                'split_files': output_files,
                'status': 'success',
                'message': f'成功分割为{len(sections)//2}个章节文件'
            }, ensure_ascii=False)
        except Exception as e:
            return json5.dumps({
                'error': str(e),
                'status': 'failed'
            })

@register_tool('markdown_editor')
class MarkdownEditor(BaseTool):
    description = '提供多种Markdown编辑操作(替换/插入/重写)。输入: 文件路径和操作列表(包含操作类型、目标位置、新内容等)。输出: 编辑后的文件路径和操作结果'
    parameters = [{
        'name': 'file_path',
        'type': 'string',
        'description': '待编辑的Markdown文件路径',
        'required': True
    }, {
        'name': 'operations',
        'type': 'array',
        'description': '编辑操作列表',
        'items': {
            'type': 'object',
            'properties': {
                'type': {
                    'type': 'string',
                    'enum': ['replace', 'insert', 'rewrite'],
                    'description': '操作类型'
                },
                'target': {
                    'type': 'string',
                    'description': '目标位置(章节标题/行号/正则表达式)'
                },
                'content': {
                    'type': 'string',
                    'description': '新内容'
                },
                'position': {
                    'type': 'string',
                    'enum': ['before', 'after'],
                    'description': '插入位置(仅insert操作需要)'
                },
                'options': {
                    'type': 'object',
                    'properties': {
                        'is_regex': {
                            'type': 'boolean',
                            'description': '是否使用正则表达式'
                        },
                        'preserve_format': {
                            'type': 'boolean',
                            'description': '是否保留原格式'
                        }
                    }
                }
            },
            'required': ['type', 'target', 'content']
        },
        'required': True
    }]

    def call(self, params: str, **kwargs) -> str:
        try:
            data = json5.loads(params)
            if 'file_path' not in data:
                raise ValueError("缺少必要参数: file_path")
            if 'operations' not in data:
                raise ValueError("缺少必要参数: operations")
            return self._process_operations(data['file_path'], data['operations'])
        except Exception as e:
            return json5.dumps({
                'error': str(e),
                'status': 'failed'
            })

    def _process_operations(self, file_path, operations):
        """处理多种编辑操作并记录修改详情"""
        try:
            # 验证文件存在
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")

            # 读取原始内容
            with open(file_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()

            modified = False
            new_content = lines.copy()
            operation_details = []

            for op in operations:
                op_type = op.get('type')
                target = op.get('target')
                content = op.get('content')
                position = op.get('position', 'after')
                options = op.get('options', {})

                # 根据操作类型处理并记录修改
                detail = {
                    'type': op_type,
                    'target': target,
                    'content': content,
                    'position': position if op_type == 'insert' else None
                }
                
                if op_type == 'replace':
                    new_content, modified_count, old_content = self._replace_content(
                        new_content, target, content, options)
                    detail['old_content'] = old_content if modified_count > 0 else None
                    detail['modified_count'] = modified_count
                    changed = modified_count > 0
                elif op_type == 'insert':
                    new_content, changed = self._insert_content(
                        new_content, target, content, position, options)
                elif op_type == 'rewrite':
                    new_content, changed, old_content = self._rewrite_content(
                        new_content, target, content, options)
                    detail['old_content'] = old_content if changed else None
                else:
                    raise ValueError(f"未知操作类型: {op_type}")

                if changed:
                    # 记录修改位置(行号)
                    for i, line in enumerate(new_content):
                        if op_type == 'replace' and target in line:
                            detail['line_number'] = i + 1
                        elif op_type == 'insert' and content in line:
                            detail['line_number'] = i + 1
                    
                    # 记录编辑操作到日志
                    from qwen_agent.log import log_edit_operation
                    log_edit_operation(
                        file_path=file_path,
                        operation_type=op_type,
                        target=target,
                        new_content=content,
                        line_number=detail.get('line_number')
                    )
                
                operation_details.append(detail)
                modified = modified or changed

            # 保存修改
            if modified:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.writelines(new_content)
                    f.flush()

            # 统计实际修改位置数量
            modified_positions = sum(
                op.get('modified_count', 0) 
                for op in operation_details 
                if op.get('type') == 'replace' and op.get('old_content') is not None
            )
            
            return json5.dumps({
                'edited_path': file_path,
                'status': 'success',
                'modified': modified,
                'modified_positions': modified_positions,
                'details': operation_details,
                'message': f'共修改了{modified_positions}处内容' if modified else '内容无变化'
            }, ensure_ascii=False)
        except Exception as e:
            return json5.dumps({
                'error': str(e),
                'status': 'failed'
            })

    def _replace_content(self, lines, target, new_content, options):
        """替换内容并返回修改位置数量和旧内容"""
        is_regex = options.get('is_regex', False)
        modified_count = 0
        old_content = None
        new_lines = []

        for line in lines:
            if is_regex:
                matches = re.findall(target, line)
                if matches:
                    modified_count += len(matches)
                    old_content = line
                    line = re.sub(target, new_content, line)
            elif target in line:
                modified_count += line.count(target)
                old_content = line
                line = line.replace(target, new_content)
            new_lines.append(line)

        return new_lines, modified_count, old_content

    def _insert_content(self, lines, target, content, position, options):
        """插入内容"""
        is_regex = options.get('is_regex', False)
        changed = False
        new_lines = []

        for i, line in enumerate(lines):
            new_lines.append(line)
            match = False
            if is_regex:
                match = bool(re.search(target, line))
            else:
                match = target in line

            if match:
                if position == 'before':
                    new_lines.insert(i, content + '\n')
                else:
                    new_lines.insert(i+1, content + '\n')
                changed = True

        return new_lines, changed

    def _rewrite_content(self, lines, target, new_content, options):
        """重写内容并返回旧内容"""
        preserve_format = options.get('preserve_format', True)
        changed = False
        old_content = None
        new_lines = []

        for line in lines:
            if target in line and preserve_format:
                old_content = line
                # 保留原格式标记(如**粗体**等)
                md_elements = re.findall(r'(\*\*.*?\*\*|_.*?_|`.*?`)', line)
                new_line = new_content
                for elem in md_elements:
                    if elem in new_line:
                        new_line = new_line.replace(elem, elem)
                line = line.replace(target, new_line)
                changed = True
            elif target in line:
                old_content = line
                line = new_content
                changed = True
            new_lines.append(line)

        return new_lines, changed, old_content

@register_tool('word_to_markdown')
class WordToMarkdownTool(BaseTool):
    """Word转Markdown工具"""
    description = '转换Word(.docx)为Markdown格式，保留标题层级和基本格式。输入: Word文件路径。输出: 生成的Markdown文件路径'
    parameters = [{
        'name': 'file_path',
        'type': 'string',
        'description': 'Word文档路径',
        'required': True
    }, {
        'name': 'output_dir',
        'type': 'string',
        'description': '输出目录',
        'required': False
    }]

    def call(self, params: str, **kwargs) -> str:
        try:
            data = json5.loads(params)
            if 'file_path' not in data:
                raise ValueError("缺少必要参数: file_path")
            output_dir = data.get('output_dir', os.getcwd())
            return self._word_to_markdown(data['file_path'], output_dir)
        except Exception as e:
            return json5.dumps({
                'error': str(e),
                'status': 'failed'
            })

    def iter_block_items(self, parent):
        """按文档原始顺序遍历段落和表格"""
        from docx.document import Document
        if isinstance(parent, Document):
            parent_elm = parent.element.body
        else:
            raise ValueError("不支持的父元素类型")
        
        for child in parent_elm.iterchildren():
            if child.tag.endswith('p'):  # 段落
                yield Paragraph(child, parent)
            elif child.tag.endswith('tbl'):  # 表格
                yield Table(child, parent)

    def _word_to_markdown(self, docx_path, output_dir):
        """将Word文档转换为Markdown格式，保留表格、图片和段落间距"""
        doc = Document(docx_path)
        md_content = ""
        image_count = 0
        
        for block in self.iter_block_items(doc):
            if isinstance(block, Paragraph):
                # 计算段落间距换行数 (1pt ≈ 0.35px, 1换行≈15px)
                space_before = int(block.paragraph_format.space_before.pt / 15) if block.paragraph_format.space_before else 0
                space_after = int(block.paragraph_format.space_after.pt / 15) if block.paragraph_format.space_after else 0
                
                # 添加段前间距
                md_content += '\n' * max(1, space_before)
                
                # 处理标题
                if block.style.name.startswith('Heading'):
                    level = int(block.style.name[-1])
                    md_content += f"{'#' * level} {block.text}\n\n"
                # 处理普通段落
                else:
                    text = block.text
                    # 保留格式标记
                    for run in block.runs:
                        if run.bold:
                            text = text.replace(run.text, f"**{run.text}**")
                        if run.italic:
                            text = text.replace(run.text, f"*{run.text}*")
                        if run.underline:
                            text = text.replace(run.text, f"<u>{run.text}</u>")
                    md_content += f"{text}\n"
                
                # 添加段后间距
                md_content += '\n' * max(1, space_after)
                
                # 处理图片
                for run in block.runs:
                    if run._element.xpath('.//pic:pic'):
                        image_count += 1
                        md_content += self._process_image(run, output_dir, image_count)
            
            elif isinstance(block, Table):
                md_content += self._convert_table_to_markdown(block) + "\n\n"
        
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        md_path = os.path.join(os.getcwd(), f"{base_name}.md")
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return json5.dumps({
            'markdown_path': md_path,
            'status': 'success',
            'message': '文件已保存到当前工作目录'
        }, ensure_ascii=False)

    def _process_image(self, run, output_dir, image_count):
        """处理图片并转换为Markdown格式"""
        try:
            os.makedirs(output_dir, exist_ok=True)
            for pic in run._element.xpath('.//pic:pic'):
                embed = pic.xpath('.//a:blip/@r:embed')[0]
                image_part = run.part.related_parts[embed]
                image_data = image_part.blob
                
                image_name = f"image_{image_count}.png"
                abs_output_dir = os.path.abspath(output_dir)
                image_path = os.path.join(abs_output_dir, image_name)
                
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                
                return f"![image {image_count}]({image_name})\n\n"
        except Exception as e:
            print(f"图片处理错误: {str(e)}")
        return ""

    def _convert_table_to_markdown(self, table):
        """将Word表格转换为Markdown表格格式"""
        # 表头
        header = "| " + " | ".join(cell.text for cell in table.rows[0].cells) + " |\n"
        # 分隔线
        separator = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |\n"
        # 表格内容
        body = ""
        for row in table.rows[1:]:
            body += "| " + " | ".join(cell.text for cell in row.cells) + " |\n"
        
        return header + separator + body + "\n"

@register_tool('markdown_to_word')
class MarkdownToWordTool(BaseTool):
    """Markdown转Word工具"""
    description = '转换Markdown为Word(.docx)格式，保留标题层级。输入: Markdown文件路径。输出: 生成的Word文件路径'
    parameters = [{
        'name': 'file_path',
        'type': 'string',
        'description': 'Markdown文件路径',
        'required': True
    }, {
        'name': 'output_dir',
        'type': 'string',
        'description': '输出目录',
        'required': False
    }]

    def call(self, params: str, **kwargs) -> str:
        try:
            data = json5.loads(params)
            if 'file_path' not in data:
                raise ValueError("缺少必要参数: file_path")
            output_dir = data.get('output_dir', os.getcwd())
            return self._markdown_to_word(data['file_path'], output_dir)
        except Exception as e:
            return json5.dumps({
                'error': str(e),
                'status': 'failed'
            })

    def _markdown_to_word(self, md_path, output_dir):
        """将Markdown转换回Word文档，保留表格和格式"""
        doc = Document()
        
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        current_para = None
        in_table = False
        table_rows = []
        
        for line in lines:
            line = line.strip()
            
            # 处理表格
            if line.startswith('|') and line.endswith('|'):
                if not in_table:
                    in_table = True
                    table_rows = []
                table_rows.append([cell.strip() for cell in line.split('|')[1:-1]])
                continue
            elif in_table:
                # 结束表格处理
                if len(table_rows) > 1:  # 至少表头+分隔线
                    table = doc.add_table(rows=1, cols=len(table_rows[0]))
                    hdr_cells = table.rows[0].cells
                    for i, cell in enumerate(table_rows[0]):
                        hdr_cells[i].text = cell
                    
                    for row in table_rows[2:]:  # 跳过表头和分隔线
                        row_cells = table.add_row().cells
                        for i, cell in enumerate(row):
                            row_cells[i].text = cell
                in_table = False
            
            # 处理标题
            if line.startswith('#'):
                level = line.count('#')
                doc.add_heading(line.lstrip('#').strip(), level=level)
            # 处理图片
            elif line.startswith('!['):
                match = re.match(r'!\[.*?\]\((.*?)\)', line)
                if match:
                    try:
                        image_path = os.path.join(os.path.dirname(md_path), match.group(1))
                        doc.add_picture(image_path)
                    except:
                        pass
            # 处理普通段落
            elif line:
                if current_para is None:
                    current_para = doc.add_paragraph(line)
                else:
                    current_para.add_run('\n' + line)
            else:
                current_para = None
        
        base_name = os.path.splitext(os.path.basename(md_path))[0]
        docx_path = os.path.join(os.getcwd(), f"{base_name}.docx")
        doc.save(docx_path)
        
        return json5.dumps({
            'word_path': docx_path,
            'status': 'success',
            'message': '文件已保存到当前工作目录'
        }, ensure_ascii=False)

@register_tool('markdown_merger')
class ChapterMarkdownMerger(BaseTool):
    """Markdown章节合并工具"""
    description = '按指定顺序合并多个Markdown文件，用分隔线区分章节。输入: Markdown文件列表。输出: 合并后的文件路径'
    parameters = [{
        'name': 'file_paths',
        'type': 'array',
        'description': '按章节顺序排列的Markdown文件列表',
        'required': True
    }, {
        'name': 'output_path',
        'type': 'string',
        'description': '输出文件路径(默认merged.md)',
        'required': False
    }]

    def call(self, params: str, **kwargs) -> str:
        try:
            data = json5.loads(params)
            if 'file_paths' not in data:
                raise ValueError("缺少必要参数: file_paths")
            output_path = data.get('output_path', 'merged.md')
            return self._merge_chapters(data['file_paths'], output_path)
        except Exception as e:
            return json5.dumps({
                'error': str(e),
                'status': 'failed'
            })

    def _merge_chapters(self, md_paths, output_path):
        """合并Markdown章节文件"""
        # 验证输入必须是列表
        if not isinstance(md_paths, list):
            raise ValueError("file_paths必须是列表")
        
        # 验证文件存在性
        valid_paths = []
        for path in md_paths:
            abs_path = os.path.abspath(path)
            if not os.path.exists(abs_path):
                raise FileNotFoundError(f"文件不存在: {abs_path}")
            if not abs_path.lower().endswith('.md'):
                raise ValueError(f"非Markdown文件: {abs_path}")
            valid_paths.append(abs_path)
        
        # 按顺序合并内容
        merged_content = ""
        for i, path in enumerate(valid_paths):
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read().strip()
                if i > 0:  # 不是第一个文件，添加分隔符
                    merged_content += '\n\n---\n\n'
                merged_content += content
        
        # 写入输出文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(merged_content)
        
        return json5.dumps({
            'merged_path': os.path.abspath(output_path),
            'status': 'success',
            'message': f'成功合并{len(valid_paths)}个章节',
            'merged_files': [os.path.basename(p) for p in valid_paths]
        }, ensure_ascii=False)

def init_agent_service():
    llm_cfg = {
        'model': 'deepseek-v3-0324',
        'model_server': 'http://172.16.56.101:8099/v1', 
        'api_key': '64a6c2dc-56bc-4a93-b9d1-5da69e3a6abe'
    }
    # 注册并初始化工具
    word_to_md = WordToMarkdownTool()
    md_to_word = MarkdownToWordTool()
    md_merger = ChapterMarkdownMerger()
    markdown_splitter = MarkdownSplitter()
    markdown_editor = MarkdownEditor()
    
    # 工具列表
    tools = [word_to_md, md_to_word, md_merger, markdown_splitter, markdown_editor]
    
    system = '''你是一个专业的文档处理专家，支持以下功能：
1. Word转Markdown格式
2. Markdown转Word格式
3. Markdown按章节拆分
4. Markdown章节内容编辑
5. 合并多个Markdown文件
6. 表格与列表格式转换
7. 图片格式保持

使用流程：
当你收到两个文件时，你将会先将初稿转换成markdown文件，然后按照章节拆分成几个文件，按照修改方案修改对应章节内容，最后修改好以后将拆分的章节合并，转换回word文件，最后返回地址或下载地址。

注意事项：
1. 所有编辑操作需保持原始格式不变
2. 回答用中文
'''
    
    bot = VirtualMemoryAgent(
        llm=llm_cfg,
        system_message=system,
        function_list=tools,
    )
    return bot

if __name__ == '__main__':
    bot = init_agent_service()
    WebUI(bot).run()
